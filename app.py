"""
app.py — نظام الأقسام: دخول ← قسم ← رفع ← استخراج (Gemini)
← تصحيح تلقائي (Arabic KB) ← تدقيق بصري (Claude: اقتراحات نصية)
← أزرار تصحيح للمراجع البشري ← Word + إكسل معبّأ.
"""
import difflib
import functools
import json
import os
import re
import shutil
import threading

from flask import (Flask, render_template, request, redirect, url_for,
                   send_file, abort, session, flash, jsonify)
from werkzeug.utils import secure_filename
from werkzeug.security import check_password_hash
from docx import Document as DocxDocument
from docx.oxml.ns import qn

import config as C
from models import (init_db, SessionLocal, User, Section, Document,
                    Page, Correction)
from pipeline import pdf_utils, process, template_utils
from arabic_kb.arabic_kb_models import (
    init_arabic_kb, log_correction, build_smart_hints
)
from arabic_kb.lexicon_models import (
    init_lexicon, add_term, add_terms_from_text, lexicon_stats
)
from arabic_kb.suggestions import (
    init_suggestions, build_for_page, apply_suggestion,
    reject_suggestion, get_page_suggestions, suggestion_accuracy
)


app = Flask(__name__)
app.secret_key = C.SECRET_KEY
app.config["MAX_CONTENT_LENGTH"] = C.MAX_UPLOAD_MB * 1024 * 1024

for d in (C.UPLOADS, C.STORAGE, C.TEMPLATES_DIR, C.OUTPUTS_DIR):
    os.makedirs(d, exist_ok=True)

# تهيئة قواعد البيانات
init_db()
init_arabic_kb()
init_lexicon()
init_suggestions()


def _rtl(par):
    pPr = par._p.get_or_add_pPr()
    pPr.append(pPr.makeelement(qn('w:bidi'), {}))
    return par


# ---------- الدخول ----------

def login_required(f):
    @functools.wraps(f)
    def wrapper(*a, **k):
        if not session.get("user"):
            return redirect(url_for("login", next=request.path))
        return f(*a, **k)
    return wrapper


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        db = SessionLocal()
        u = db.query(User).filter_by(
            username=request.form.get("username", "")).first()
        db.close()
        if u and check_password_hash(u.password_hash,
                                     request.form.get("password", "")):
            session["user"] = u.username
            return redirect(request.args.get("next") or url_for("sections"))
        flash("اسم المستخدم أو كلمة المرور غير صحيحة")
    return render_template("login.html")


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))


# ---------- الأقسام ----------

@app.route("/")
@login_required
def home():
    return redirect(url_for("sections"))


@app.route("/sections")
@login_required
def sections():
    db = SessionLocal()
    secs = db.query(Section).order_by(Section.id.desc()).all()
    counts = {s.id: len(s.documents) for s in secs}
    db.close()
    return render_template("sections.html", sections=secs, counts=counts,
                           mock=C.MOCK, engine=C.STRUCTURE_ENGINE)


@app.route("/sections/create", methods=["POST"])
@login_required
def create_section():
    name = request.form.get("name", "").strip()
    desc = request.form.get("description", "").strip()
    tpl  = request.files.get("template")
    has_tpl = tpl and tpl.filename

    if not name:
        flash("أدخل اسم القسم")
        return redirect(url_for("sections"))
    if has_tpl and not tpl.filename.lower().endswith((".xlsx", ".xlsm")):
        flash("القالب يجب أن يكون ملف إكسل (.xlsx)")
        return redirect(url_for("sections"))

    db = SessionLocal()
    sec = Section(name=name, description=desc)
    db.add(sec)
    db.commit()

    if has_tpl:
        tpl_path = os.path.join(C.TEMPLATES_DIR, f"section_{sec.id}.xlsx")
        tpl.save(tpl_path)
        try:
            fields = template_utils.parse_template(tpl_path)
        except Exception as e:
            db.delete(sec)
            db.commit()
            db.close()
            flash(f"تعذّرت قراءة القالب: {e}")
            return redirect(url_for("sections"))
        sec.template_path = tpl_path
        sec.fields_json   = json.dumps(fields, ensure_ascii=False)
        db.commit()

    sid = sec.id
    db.close()
    return redirect(url_for("section_home", section_id=sid))


@app.route("/section/<int:section_id>")
@login_required
def section_home(section_id):
    db  = SessionLocal()
    sec = db.get(Section, section_id)
    if not sec:
        db.close()
        abort(404)
    fields  = json.loads(sec.fields_json) if sec.fields_json else {}
    learned = (db.query(Correction).filter_by(section_id=sec.id)
               .order_by(Correction.count.desc()).limit(10).all())
    data = {
        "id": sec.id, "name": sec.name, "description": sec.description,
        "learned": [{"label": c.label, "wrong": c.wrong,
                     "right": c.right, "count": c.count} for c in learned],
        "has_template": bool(sec.template_path),
        "kv_count":   len(fields.get("kv", [])),
        "table_cols": (fields.get("table") or {}).get("columns", []),
        "docs": [{"id": d.id, "filename": d.filename, "pages": d.total_pages,
                  "status": d.status, "score": d.score,
                  "has_output": bool(d.output_path)}
                 for d in sec.documents],
    }
    db.close()
    return render_template("section.html", s=data)


# ---------- رفع ومعالجة ----------

@app.route("/section/<int:section_id>/upload", methods=["POST"])
@login_required
def upload(section_id):
    db  = SessionLocal()
    sec = db.get(Section, section_id)
    if not sec:
        db.close()
        abort(404)
    f = request.files.get("file")
    if not f or not f.filename:
        db.close()
        return redirect(url_for("section_home", section_id=section_id))

    fname = secure_filename(f.filename) or "document.pdf"
    saved = os.path.join(C.UPLOADS, fname)
    f.save(saved)

    doc = Document(filename=fname, section_id=sec.id, status="processing")
    db.add(doc)
    db.commit()

    page_dir = os.path.join(C.STORAGE, str(doc.id))
    try:
        images = pdf_utils.split_to_images(saved, page_dir, dpi=300)
    except Exception as e:
        doc.status = "error"
        db.commit()
        db.close()
        return f"خطأ في تقسيم الملف: {e}", 400

    doc.total_pages = len(images)
    db.commit()
    doc_id = doc.id
    db.close()

    threading.Thread(target=_process_document,
                     args=(doc_id, images, fname), daemon=True).start()
    return redirect(url_for("detail", doc_id=doc_id))


def _process_document(doc_id, images, fname):
    """يعالج الوثيقة صفحة صفحة في الخلفية."""
    db  = SessionLocal()
    doc = db.get(Document, doc_id)
    sec = doc.section

    hints = build_smart_hints(sec.id) if sec else ""

    try:
        corrected_texts, scores, all_notes = [], [], []

        for i, img in enumerate(images, 1):
            out = process.process_page(img, hints=hints)
            corrected_texts.append(f"=== صفحة {i} ===\n{out['corrected_text']}")
            if out["score"] is not None:
                scores.append(out["score"])
            for n in out["notes"]:
                all_notes.append(f"صفحة {i}: {n}")

            page_obj = Page(
                document_id=doc.id, page_no=i, image_path=img,
                raw_text=out["raw_text"],
                corrected_text=out["corrected_text"],
                score=out["score"],
                notes=json.dumps(out["notes"], ensure_ascii=False),
                structured_json=json.dumps(out["result"], ensure_ascii=False),
                needs_review=out["needs_review"],
                status="done",
            )
            db.add(page_obj)
            db.commit()

            # بناء اقتراحات الصفحة: قاموس محلي + Claude
            try:
                build_for_page(page_obj.id,
                               out["corrected_text"],
                               out.get("suggestions", []))
            except Exception as e:
                print(f"تعذّر بناء اقتراحات الصفحة {i}: {e}")

        full_text = "\n\n".join(corrected_texts)
        values    = {}

        if sec and sec.fields_json:
            fields     = json.loads(sec.fields_json)
            kv_labels, table_cols = template_utils.labels_of(fields)
            values = process.structure_for_section(
                kv_labels, table_cols, full_text, hints=hints)
            out_name = f"doc_{doc.id}_{os.path.splitext(fname)[0]}.xlsx"
            out_path = os.path.join(C.OUTPUTS_DIR, out_name)
            template_utils.fill_template(
                sec.template_path, fields, values, out_path)
            doc.output_path = out_path

        tpl_review = values.get("مراجعة", []) if values else []
        all_review = all_notes + [f"حقل: {r}" for r in tpl_review]
        doc.merged_json  = json.dumps(
            {**values, "ملاحظات_التدقيق": all_review}, ensure_ascii=False)
        doc.needs_review = bool(all_review)
        doc.score        = round(sum(scores) / len(scores), 1) if scores else None
        doc.status       = "done"
        db.commit()

    except Exception as e:
        db.rollback()
        try:
            doc = db.get(Document, doc_id)
            doc.status      = "error"
            doc.merged_json = json.dumps({"خطأ": str(e)}, ensure_ascii=False)
            db.commit()
        except Exception:
            db.rollback()
    finally:
        db.close()


# ---------- ذاكرة التعلم ----------

def _save_correction(db, section_id, kind, label, wrong, right):
    wrong, right = (wrong or "").strip(), (right or "").strip()
    if not wrong or not right or wrong == right:
        return
    if len(wrong) > 80 or len(right) > 80:
        return
    row = (db.query(Correction)
           .filter_by(section_id=section_id, kind=kind,
                      label=label, wrong=wrong, right=right).first())
    if row:
        row.count += 1
    else:
        db.add(Correction(section_id=section_id, kind=kind,
                          label=label, wrong=wrong, right=right))


def _harvest_text_diff(db, section_id, old_text, new_text, limit=10):
    ow, nw = (old_text or "").split(), (new_text or "").split()
    sm     = difflib.SequenceMatcher(a=ow, b=nw, autojunk=False)
    saved  = 0
    for op, i1, i2, j1, j2 in sm.get_opcodes():
        if op != "replace" or saved >= limit:
            continue
        wrong, right = " ".join(ow[i1:i2]), " ".join(nw[j1:j2])
        if 0 < len(wrong) <= 60 and 0 < len(right) <= 60:
            _save_correction(db, section_id, "text", "", wrong, right)
            saved += 1


def _build_section_hints(db, section_id, max_items=15):
    """يجمع المعرفة الثابتة + تصحيحات القسم السابقة."""
    from arabic_kb.arabic_kb_models import OCRCorrectionLog
    from arabic_kb.arabic_patterns import build_ocr_hints_prompt

    lines = [build_ocr_hints_prompt()]

    old_rows = (db.query(Correction).filter_by(section_id=section_id)
                .order_by(Correction.count.desc(),
                          Correction.updated_at.desc())
                .limit(max_items).all())

    new_rows = (db.query(OCRCorrectionLog).filter_by(section_id=section_id)
                .order_by(OCRCorrectionLog.occurrence_count.desc())
                .limit(max_items).all())

    if old_rows or new_rows:
        lines.append("\n【تصحيحات بشرية موثوقة لهذا القسم — طبّقها فوراً】")
        for r in old_rows:
            prefix = f"في حقل «{r.label}»: " if r.kind == "field" and r.label else ""
            lines.append(f"• {prefix}«{r.wrong}» → «{r.right}» [{r.count}x]")
        for c in new_rows:
            ctx   = f" (…{c.context_before}…)" if c.context_before else ""
            field = f" في «{c.field_label}»"   if c.field_label    else ""
            lines.append(
                f"• [{c.correction_type}]{field}: "
                f"«{c.ocr_reading}» → «{c.human_correction}»"
                f"{ctx} [{c.occurrence_count}x]"
            )

    return "\n".join(lines)


# ---------- عرض وثيقة ----------

@app.route("/document/<int:doc_id>/progress")
@login_required
def progress(doc_id):
    db  = SessionLocal()
    doc = db.get(Document, doc_id)
    if not doc:
        db.close()
        abort(404)
    done = db.query(Page).filter(Page.document_id == doc_id).count()
    data = {"status": doc.status, "total": doc.total_pages or 0,
            "done": done, "score": doc.score}
    db.close()
    return jsonify(data)


@app.route("/document/<int:doc_id>")
@login_required
def detail(doc_id):
    db  = SessionLocal()
    doc = db.get(Document, doc_id)
    if not doc:
        db.close()
        abort(404)
    values = json.loads(doc.merged_json) if doc.merged_json else {}
    pages  = [{"no": p.page_no, "id": p.id, "raw": p.raw_text,
               "corrected": p.corrected_text, "score": p.score,
               "notes": json.loads(p.notes) if p.notes else [],
               "suggestions": get_page_suggestions(p.id)}
              for p in doc.pages]
    data = {
        "id": doc.id, "filename": doc.filename, "status": doc.status,
        "section_id":   doc.section_id,
        "section_name": doc.section.name if doc.section else "",
        "has_output":   bool(doc.output_path),
        "score":        doc.score,
        "filled":       values.get("حقول", {}),
        "table":        values.get("جدول", []),
        "review":       values.get("ملاحظات_التدقيق", []),
        "error":        values.get("خطأ"),
        "pages":        pages,
    }
    db.close()
    return render_template("detail.html", d=data)


@app.route("/document/<int:doc_id>/download")
@login_required
def download(doc_id):
    db  = SessionLocal()
    doc = db.get(Document, doc_id)
    db.close()
    if not doc or not doc.output_path or not os.path.exists(doc.output_path):
        abort(404)
    return send_file(doc.output_path, as_attachment=True,
                     download_name=os.path.basename(doc.output_path))


@app.route("/document/<int:doc_id>/download_word")
@login_required
def download_word(doc_id):
    db  = SessionLocal()
    doc = db.get(Document, doc_id)
    if not doc:
        db.close()
        abort(404)
    values    = json.loads(doc.merged_json) if doc.merged_json else {}
    pages     = [(p.page_no, p.corrected_text or p.raw_text, p.score)
                 for p in doc.pages]
    base_name = os.path.splitext(doc.filename)[0]
    db.close()

    w = DocxDocument()
    _rtl(w.add_heading(base_name, level=1))

    filled = values.get("حقول", {})
    if filled:
        _rtl(w.add_heading("الحقول المستخرجة", level=2))
        for k, v in filled.items():
            _rtl(w.add_paragraph(f"{k}: {v}"))

    table_rows = values.get("جدول", [])
    if table_rows:
        _rtl(w.add_heading("الجدول", level=2))
        for row in table_rows:
            _rtl(w.add_paragraph(" | ".join(str(c) for c in row if c)))

    for no, text, score in pages:
        _rtl(w.add_heading(f"صفحة {no}", level=2))
        for line in (text or "").splitlines():
            _rtl(w.add_paragraph(line))

    out = os.path.join(C.OUTPUTS_DIR, f"doc_{doc_id}_{base_name}.docx")
    w.save(out)
    return send_file(out, as_attachment=True,
                     download_name=os.path.basename(out))


# ---------- الاقتراحات ----------

@app.route("/page/<int:page_id>/suggestions")
@login_required
def page_suggestions(page_id):
    """يجلب اقتراحات الصفحة المعلّقة."""
    return jsonify({"suggestions": get_page_suggestions(page_id)})


@app.route("/page/<int:page_id>/apply/<int:sug_id>", methods=["POST"])
@login_required
def apply_sug(page_id, sug_id):
    """يطبّق اقتراحاً على نص الصفحة ويحفظه."""
    db = SessionLocal()
    try:
        p = db.get(Page, page_id)
        if not p:
            abort(404)
        # النص المرسل من الواجهة له الأولوية (قد يكون المستخدم عدّل يدوياً)
        current = request.form.get("text") or p.corrected_text or ""
        res = apply_suggestion(sug_id, current)
        if res["ok"]:
            p.corrected_text = res["text"]
            db.commit()
        return jsonify(res)
    finally:
        db.close()


@app.route("/page/<int:page_id>/reject/<int:sug_id>", methods=["POST"])
@login_required
def reject_sug(page_id, sug_id):
    """يرفض اقتراحاً — والكلمة الأصلية تدخل القاموس كصحيحة."""
    return jsonify(reject_suggestion(sug_id))


@app.route("/lexicon/stats")
@login_required
def lex_stats():
    """إحصاءات القاموس ودقة الاقتراحات حسب المصدر."""
    return jsonify({"lexicon": lexicon_stats(),
                    "accuracy": suggestion_accuracy()})


@app.route("/lexicon/add", methods=["POST"])
@login_required
def lex_add():
    """إضافة مصطلح يدوياً للقاموس."""
    term  = request.form.get("term", "").strip()
    ttype = request.form.get("type", "person")
    ok = add_term(term, ttype, source="verified")
    return jsonify({"ok": ok, "term": term})


# ---------- تعديل وتدقيق بعد المعالجة ----------

@app.route("/document/<int:doc_id>/update_fields", methods=["POST"])
@login_required
def update_fields(doc_id):
    db  = SessionLocal()
    doc = db.get(Document, doc_id)
    if not doc:
        db.close()
        abort(404)
    values     = json.loads(doc.merged_json) if doc.merged_json else {}
    old_fields = values.get("حقول", {})
    new_fields = {k: v.strip() for k, v in request.form.items()}

    if doc.section_id:
        for k, new_v in new_fields.items():
            old_v = old_fields.get(k, "")
            _save_correction(db, doc.section_id, "field", k, old_v, new_v)
            if old_v and new_v and old_v != new_v:
                if "تاريخ" in k:
                    ctype = "date"
                elif "وطني" in k or "هوية" in k:
                    ctype = "national_id"
                elif "مبلغ" in k or "سعر" in k:
                    ctype = "amount"
                elif "اسم" in k:
                    ctype = "name"
                elif re.search(r"[٠-٩0-9]", old_v):
                    ctype = "digit"
                else:
                    ctype = "other"
                log_correction(
                    section_id=doc.section_id,
                    page_id=None,
                    ocr_reading=old_v,
                    human_correction=new_v,
                    correction_type=ctype,
                    field_label=k,
                )
                # اسم أو منطقة مصحَّحة يدوياً ← القاموس
                if ctype == "name":
                    add_term(new_v, "person", source="verified")
                elif "منطقة" in k or "المنطقة" in k:
                    add_term(new_v, "region", source="verified")

    values["حقول"] = new_fields
    sec = doc.section
    if sec and sec.fields_json and sec.template_path and doc.output_path:
        fields = json.loads(sec.fields_json)
        template_utils.fill_template(
            sec.template_path, fields, values, doc.output_path)

    doc.merged_json = json.dumps(values, ensure_ascii=False)
    db.commit()
    db.close()
    flash("حُفظت التعديلات وتم تحديث ملف الإكسل")
    return redirect(url_for("detail", doc_id=doc_id))


@app.route("/page/<int:page_id>/update_text", methods=["POST"])
@login_required
def update_text(page_id):
    db = SessionLocal()
    p  = db.get(Page, page_id)
    if not p:
        db.close()
        abort(404)
    new_text = request.form.get("text", "")
    old_text = p.corrected_text or ""

    if p.document and p.document.section_id:
        sid = p.document.section_id

        _harvest_text_diff(db, sid, old_text, new_text)

        ow, nw = old_text.split(), new_text.split()
        sm = difflib.SequenceMatcher(a=ow, b=nw, autojunk=False)
        for op, i1, i2, j1, j2 in sm.get_opcodes():
            if op != "replace":
                continue
            wrong = " ".join(ow[i1:i2])
            right = " ".join(nw[j1:j2])
            if not wrong or not right or wrong == right:
                continue
            ctx_before = " ".join(ow[max(0, i1 - 3):i1])
            ctx_after  = " ".join(ow[i2:min(len(ow), i2 + 3)])
            if re.search(r"[٠-٩0-9]", wrong):
                ctype = "digit"
            else:
                ctype = "word"
            log_correction(
                section_id=sid,
                page_id=page_id,
                ocr_reading=wrong,
                human_correction=right,
                correction_type=ctype,
                context_before=ctx_before,
                context_after=ctx_after,
            )

    p.corrected_text = new_text
    doc_id = p.document_id
    db.commit()
    db.close()

    # التعديل اليدوي مصدر موثوق — يغذّي القاموس
    try:
        add_terms_from_text(new_text)
    except Exception as e:
        print(f"تعذّر تغذية القاموس: {e}")

    flash("حُفظ نص الصفحة")
    return redirect(url_for("detail", doc_id=doc_id))


@app.route("/document/<int:doc_id>/approve", methods=["POST"])
@login_required
def approve(doc_id):
    """اعتماد الوثيقة — النص المعتمد بشرياً أوثق مصدر للقاموس."""
    db  = SessionLocal()
    doc = db.get(Document, doc_id)
    texts = []
    if doc:
        texts = [p.corrected_text for p in doc.pages if p.corrected_text]
        doc.needs_review = False
        doc.status       = "approved"
        db.commit()
    db.close()

    added = 0
    for t in texts:
        try:
            added += add_terms_from_text(t)
        except Exception as e:
            print(f"تعذّرت تغذية القاموس: {e}")

    flash(f"اعتُمدت الوثيقة — أُضيف {added} مصطلح للقاموس")
    return redirect(url_for("detail", doc_id=doc_id))


@app.route("/document/<int:doc_id>/delete", methods=["POST"])
@login_required
def delete_document(doc_id):
    db  = SessionLocal()
    doc = db.get(Document, doc_id)
    if not doc:
        db.close()
        abort(404)
    sid = doc.section_id
    _delete_document_files(doc)
    db.delete(doc)
    db.commit()
    db.close()
    flash("حُذفت الوثيقة وملفاتها")
    return redirect(url_for("section_home", section_id=sid)
                    if sid else url_for("sections"))


@app.route("/section/<int:section_id>/delete", methods=["POST"])
@login_required
def delete_section(section_id):
    db  = SessionLocal()
    sec = db.get(Section, section_id)
    if not sec:
        db.close()
        abort(404)
    for d in list(sec.documents):
        _delete_document_files(d)
        db.delete(d)
    if sec.template_path and os.path.exists(sec.template_path):
        os.remove(sec.template_path)
    db.delete(sec)
    db.commit()
    db.close()
    flash("حُذف القسم وكل وثائقه")
    return redirect(url_for("sections"))


def _delete_document_files(doc):
    page_dir = os.path.join(C.STORAGE, str(doc.id))
    if os.path.isdir(page_dir):
        shutil.rmtree(page_dir, ignore_errors=True)
    if doc.output_path and os.path.exists(doc.output_path):
        os.remove(doc.output_path)
    base  = os.path.splitext(doc.filename)[0]
    docx  = os.path.join(C.OUTPUTS_DIR, f"doc_{doc.id}_{base}.docx")
    if os.path.exists(docx):
        os.remove(docx)


@app.route("/page/<int:page_id>/image")
@login_required
def page_image(page_id):
    db = SessionLocal()
    p  = db.get(Page, page_id)
    db.close()
    if not p or not os.path.exists(p.image_path):
        abort(404)
    return send_file(p.image_path)


# ---------- بحث ----------

@app.route("/search")
@login_required
def search():
    q       = request.args.get("q", "").strip()
    results = []
    if q:
        db   = SessionLocal()
        like = f"%{q}%"
        pages = db.query(Page).filter(
            (Page.corrected_text.like(like)) |
            (Page.raw_text.like(like))
        ).all()
        results = [{"doc_id": p.document_id, "no": p.page_no} for p in pages]
        db.close()
    return render_template("search.html", q=q, results=results)


if __name__ == "__main__":
    app.run(host="0.0.0.0", debug=True, port=5000)